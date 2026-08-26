from pathlib import Path

from md2style.engines.learner import Learner
from md2style.engines import StyleEngine


def test_learn_and_reuse(tmp_path):
    # 构造一个带样式的 docx
    from docx import Document
    doc = Document()
    h = doc.add_heading("标题", level=1)
    h.runs[0].font.size = __import__("docx").shared.Pt(24)
    p = doc.add_paragraph("正文")
    p.runs[0].font.size = __import__("docx").shared.Pt(12)
    tpl = tmp_path / "tpl.docx"
    doc.save(str(tpl))

    styles_dir = tmp_path / "styles"
    styles_dir.mkdir()
    learner = Learner(styles_dir)
    yaml_path = learner.learn(str(tpl), "my_learned")

    # 闭环：学习所得应能被 StyleEngine 加载校验
    se = StyleEngine(styles_dir)
    final = se.resolve("my_learned", {})
    assert final["headings"]["h1"]["size"] == 24
