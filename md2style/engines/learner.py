"""Learner：从 Word 版式反推声明式 YAML 的提取器（数据流转规则层，逆向闭环）。

双模式：
- 模式A（精准）：读标准样式 Heading 1/2、Normal 直接映射。
- 模式B（启发式）：按字号/加粗聚类 → H1/H2/body；主色取首个非黑颜色。
输出与 StyleConfig 同构，保证 learn 产物可被 Style Engine 复用。
"""

import re
from pathlib import Path

import yaml
from docx import Document

from ..core.errors import Md2StyleError
from ..utils.color import hex_to_rgb
from ..utils.logger import get_logger

logger = get_logger("learner")


class Learner:
    def __init__(self, styles_dir: str | Path = "styles"):
        self.styles_dir = Path(styles_dir)

    def learn(self, docx_path: str, name: str) -> str:
        """从模板学习样式，写入 styles/<name>.yaml，返回路径。

        支持格式：.docx（Word 逆向提取）、.html/.htm（CSS 声明式提取）。
        """
        suffix = Path(docx_path).suffix.lower()
        if suffix in (".html", ".htm"):
            config = self._extract_html(docx_path)
        elif suffix == ".docx":
            doc = Document(docx_path)
            config = self._extract(doc)
        else:
            raise Md2StyleError(f"学习模式不支持的格式: {suffix}（支持 docx/html）")
        out = self.styles_dir / f"{name}.yaml"
        header = f"# 自动学习自: {Path(docx_path).name}\n"
        with out.open("w", encoding="utf-8") as f:
            f.write(header)
            yaml.safe_dump(config, f, allow_unicode=True, sort_keys=False)
        logger.info("已学习样式 %s -> %s", name, out)
        return str(out)

    # ---- 提取（docx）----
    def _extract(self, doc: Document) -> dict:
        try:
            return self._precise(doc)
        except Exception as e:  # 兜底启发式
            logger.warning("精准模式失败(%s)，转启发式", e)
            return self._heuristic(doc)

    # ---- 提取（html）----
    def _extract_html(self, html_path: str) -> dict:
        """从 HTML 的声明式样式（内联 style + <style> 选择器/CSS 变量）提取。

        与 StyleConfig 同构；缺失项回退 DEFAULT_STYLE。
        """
        from ..core.constants import DEFAULT_STYLE
        from ._html_style_parser import extract_html_style

        raw = Path(html_path).read_text(encoding="utf-8", errors="ignore")
        extracted = extract_html_style(raw)  # {headings:{h1..}, body:{}, extras:{}}

        config = {
            "headings": {},
            "body": dict(DEFAULT_STYLE["body"]),
            "code": dict(DEFAULT_STYLE["code"]),
        }
        for level in range(1, 7):
            key = f"h{level}"
            src = extracted["headings"].get(key) or {}
            base = DEFAULT_STYLE["headings"][key]
            config["headings"][key] = {
                "size": src.get("size", base["size"]),
                "color": src.get("color", base["color"]),
                "font": src.get("font", base["font"]),
            }
        if extracted["body"]:
            b = extracted["body"]
            for fld in ("font", "size", "line_height", "color"):
                if fld in b:
                    config["body"][fld] = b[fld]
        # 把 html 专属变量（divider/accent/card 等）作为扩展字段保留
        for k, v in extracted.get("extras", {}).items():
            config[k] = v
        return config

    def _precise(self, doc: Document) -> dict:
        styles = doc.styles

        def font_of(style, sample_run=None):
            # 优先取段落 run 实际字号/颜色（所见版式），回退到样式属性
            f = style.font
            size = None
            color = None
            if sample_run is not None and sample_run.font:
                if sample_run.font.size:
                    size = int(sample_run.font.size.pt)
                if sample_run.font.color and sample_run.font.color.rgb:
                    color = "#%02X%02X%02X" % _rgb(sample_run.font.color.rgb)
            if size is None:
                size = int(f.size.pt) if f.size else 12
            if color is None:
                color = "#%02X%02X%02X" % _rgb(f.color.rgb) if f.color and f.color.rgb else "#000000"
            font_name = (sample_run.font.name if sample_run and sample_run.font.name
                         else f.name) or "微软雅黑"
            return {"size": size, "color": color, "font": font_name}

        h1_run = self._first_run_of_heading(doc, 1)
        h2_run = self._first_run_of_heading(doc, 2)
        body_run = self._first_body_run(doc)

        h1_style = doc.styles["Heading 1"] if "Heading 1" in doc.styles else None
        h2_style = doc.styles["Heading 2"] if "Heading 2" in doc.styles else None
        normal_style = doc.styles["Normal"] if "Normal" in doc.styles else None

        config = {
            "headings": {
                "h1": font_of(h1_style, h1_run) if h1_style else {"size": 22, "color": "#000000", "font": "微软雅黑"},
                "h2": font_of(h2_style, h2_run) if h2_style else {"size": 18, "color": "#000000", "font": "微软雅黑"},
                "h3": {"size": 15, "color": "#000000", "font": "微软雅黑"},
                "h4": {"size": 13, "color": "#000000", "font": "微软雅黑"},
                "h5": {"size": 12, "color": "#000000", "font": "微软雅黑"},
                "h6": {"size": 11, "color": "#000000", "font": "微软雅黑"},
            },
            "body": font_of(normal_style, body_run) if normal_style else {"font": "微软雅黑", "size": 10.5, "line_height": 1.5, "color": "#000000"},
            "code": {"font": "Consolas", "size": 10, "color": "#333333", "background": "#F5F5F5"},
        }
        return config

    @staticmethod
    def _first_run_of_heading(doc: Document, level: int):
        for p in doc.paragraphs:
            if p.style.name == f"Heading {level}" and p.runs:
                return p.runs[0]
        return None

    @staticmethod
    def _first_body_run(doc: Document):
        for p in doc.paragraphs:
            if p.style.name == "Normal" and p.runs:
                return p.runs[0]
        # 回退：首个有 run 的非标题段落
        for p in doc.paragraphs:
            if p.runs and not p.style.name.startswith("Heading"):
                return p.runs[0]
        return None

    def _heuristic(self, doc: Document) -> dict:
        sizes: list[float] = []
        for p in doc.paragraphs:
            if p.runs and p.runs[0].font.size:
                sizes.append(p.runs[0].font.size.pt)
        if not sizes:
            raise Md2StyleError("文档无可用字号信息，无法启发式学习")
        uniq = sorted(set(sizes), reverse=True)
        h1_size = uniq[0] if len(uniq) > 0 else 22
        h2_size = uniq[1] if len(uniq) > 1 else 18
        body_size = uniq[-1]

        # 主色：首个非黑颜色
        accent = "#000000"
        for p in doc.paragraphs:
            for r in p.runs:
                if r.font.color and r.font.color.rgb:
                    rgb = _rgb(r.font.color.rgb)
                    if rgb != (0, 0, 0):
                        accent = "#%02X%02X%02X" % rgb
                        break
            if accent != "#000000":
                break

        config = {
            "headings": {
                "h1": {"size": int(h1_size), "color": accent, "font": "微软雅黑"},
                "h2": {"size": int(h2_size), "color": accent, "font": "微软雅黑"},
                "h3": {"size": int(h2_size) - 2, "color": accent, "font": "微软雅黑"},
                "h4": {"size": int(body_size) + 2, "color": "#000000", "font": "微软雅黑"},
                "h5": {"size": int(body_size) + 1, "color": "#000000", "font": "微软雅黑"},
                "h6": {"size": int(body_size), "color": "#000000", "font": "微软雅黑"},
            },
            "body": {"font": "微软雅黑", "size": int(body_size), "line_height": 1.5, "color": "#000000"},
            "code": {"font": "Consolas", "size": 10, "color": "#333333", "background": "#F5F5F5"},
        }
        return config


def _rgb(c):
    """docx RGBColor -> (r,g,b)，兼容 None。"""
    try:
        return (c[0], c[1], c[2])
    except Exception:
        return (0, 0, 0)
