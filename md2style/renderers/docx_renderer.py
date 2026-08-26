"""Docx Renderer：段落/字符级样式注入器（数据接口执行层，第2层递归）。

基于 python-docx。关键设计：
- 标题用 add_paragraph + 手动设 run 字体/对齐，而非内置 Heading 样式，
  确保 YAML 中定义的字体/颜色/字号真正生效（修复此前被内置样式覆盖的问题）。
- 支持 YAML 扩展字段：align(标题对齐) / first_indent(正文首行缩进字符) / margins(页边距cm)。
- 公文风格：标题居中、红头色、正文首行缩进2字、固定页边距。
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .base import RendererBase, RendererDispatcher
from ..core.ir import IRNode, NodeType
from ..utils.color import hex_to_rgb


class DocxRenderer(RendererBase):
    def __init__(self, template_dir: str | Path = "templates/word"):
        self.template_dir = Path(template_dir)

    def render(self, ir, final_style: dict, output_path: str) -> None:
        doc = self._load_template(final_style)
        headings = final_style.get("headings", {})
        body = final_style.get("body", {})
        docx_cfg = final_style.get("_docx", {})
        self._code_cfg = final_style.get("code", {})

        self._apply_margins(doc, docx_cfg.get("margins"))

        for node in ir:
            if node.type == NodeType.HEADING:
                h = headings.get(f"h{node.level}", {})
                align = h.get("align", docx_cfg.get("heading_align", "left"))
                p = doc.add_paragraph()
                p.alignment = self._align(align)
                self._add_runs(p, node, h)
                p.space_after = Pt(6)
            elif node.type == NodeType.PARAGRAPH:
                p = doc.add_paragraph()
                self._set_indent(p, docx_cfg.get("first_indent", 0))
                self._set_spacing(p, body.get("para_spacing", 1.0))
                self._add_runs(p, node, body)
            elif node.type == NodeType.LIST:
                p = doc.add_paragraph(style="List Bullet" if not node.ordered else "List Number")
                self._set_indent(p, docx_cfg.get("first_indent", 0))
                self._add_runs(p, node, body)
            elif node.type == NodeType.CODE:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.5)
                self._add_runs(p, node, final_style.get("code", {}))
            elif node.type == NodeType.QUOTE:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6)
                self._add_runs(p, node, body)
            elif node.type == NodeType.HR:
                doc.add_paragraph("─" * 20)
            elif node.type == NodeType.TABLE:
                self._build_table(doc, node, body, docx_cfg)

        doc.save(output_path)

    # ---- 样式 ----
    def _add_runs(self, para, node: IRNode, base_cfg: dict) -> None:
        """按结构化 runs 添加字符，正确处理加粗/斜体/行内代码/链接。"""
        # 块级字体覆盖：!{font:XXX}
        if node.meta.get("font"):
            base_cfg = dict(base_cfg, font=node.meta["font"])
        runs = node.runs
        if not runs:
            # 兜底：无结构化 runs 时用纯文本
            run = para.add_run(node.text)
            self._apply_run(run, base_cfg)
            return
        for r in runs:
            run = para.add_run(r.text)
            cfg = self._code_cfg if r.code else base_cfg
            self._apply_run(run, cfg)
            if r.bold:
                run.font.bold = True
            if r.italic:
                run.font.italic = True
            if r.link:
                run.font.underline = True
                run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xE6)

    def _apply_run(self, run, cfg: dict) -> None:
        if "font" in cfg:
            font_name = _pick_docx_font(cfg["font"])
            run.font.name = font_name
            # 同时设置东亚字体，避免中文回退到宋体
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.find(qn("w:rFonts"))
            if rfonts is None:
                rfonts = OxmlElement("w:rFonts")
                rpr.append(rfonts)
            rfonts.set(qn("w:ascii"), font_name)
            rfonts.set(qn("w:hAnsi"), font_name)
            rfonts.set(qn("w:eastAsia"), font_name)
        if "size" in cfg:
            run.font.size = Pt(cfg["size"])
        if "color" in cfg:
            r, g, b = hex_to_rgb(cfg["color"])
            run.font.color.rgb = RGBColor(r, g, b)
        if cfg.get("bold"):
            run.font.bold = True

    def _set_indent(self, para, indent_chars: int) -> None:
        if indent_chars:
            # 首行缩进：字符数 * 字号（近似）。用 first_line_chars 更准确
            para.paragraph_format.first_line_indent = Pt(indent_chars * 10.5)

    def _set_spacing(self, para, spacing_em: float) -> None:
        # 微调段落间距（em 近似为磅：1em ≈ 12pt）
        pt = Pt(max(0.0, spacing_em) * 12)
        para.paragraph_format.space_after = pt
        para.paragraph_format.space_before = pt

    def _apply_margins(self, doc, margins: dict | None) -> None:
        if not margins:
            return
        for side, cm in margins.items():
            val = Cm(cm)
            sec = doc.sections[0]
            if side == "top":
                sec.top_margin = val
            elif side == "bottom":
                sec.bottom_margin = val
            elif side == "left":
                sec.left_margin = val
            elif side == "right":
                sec.right_margin = val

    def _build_table(self, doc, node: IRNode, body: dict, docx_cfg: dict) -> None:
        rows = node.rows
        if not rows:
            return
        table = doc.add_table(rows=len(rows), cols=len(rows[0]))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"
        for r, row in enumerate(rows):
            for c, cell_runs in enumerate(row):
                cell = table.cell(r, c)
                cell_para = cell.paragraphs[0]
                if not cell_runs:
                    continue
                # 复用段落的行内渲染（保留加粗/斜体/代码/链接），\n 转为换行
                self._add_runs_list(cell_para, cell_runs, body)

    def _add_runs_list(self, para, runs: list, base_cfg: dict) -> None:
        """把结构化 runs 加入段落，处理 \n 为换行（add_break）。"""
        for r in runs:
            if r.text == "\n":
                para.add_run().add_break()
                continue
            run = para.add_run(r.text)
            cfg = self._code_cfg if r.code else base_cfg
            self._apply_run(run, cfg)
            if r.bold:
                run.font.bold = True
            if r.italic:
                run.font.italic = True
            if r.link:
                run.font.underline = True
                run.font.color.rgb = RGBColor(0x1A, 0x5C, 0xE6)

    def _align(self, align: str):
        return {
            "center": WD_ALIGN_PARAGRAPH.CENTER,
            "right": WD_ALIGN_PARAGRAPH.RIGHT,
            "left": WD_ALIGN_PARAGRAPH.LEFT,
        }.get(align, WD_ALIGN_PARAGRAPH.LEFT)

    def _load_template(self, final_style: dict):
        template_name = final_style.get("_template", "base.dotx")
        tp = self.template_dir / template_name
        if tp.exists():
            return Document(str(tp))
        return Document()


RendererDispatcher.register(".docx", DocxRenderer)


# Word 不识别字体栈（如 "-apple-system, ..."），取第一个可用的真实字体名
_SYSTEM_KEYWORDS = {"-apple-system", "blinkmacsystemfont", "segoe ui",
                    "pingfang sc", "sf pro display", "sf pro text", "system-ui"}


def _pick_docx_font(font_stack: str) -> str:
    for part in font_stack.split(","):
        name = part.strip().strip("'\"")
        if name.lower() in _SYSTEM_KEYWORDS:
            continue
        if name:  # 取首个非系统关键字字体
            return name
    return "宋体"  # 兜底
